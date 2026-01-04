import yaml
import requests
import os
import re
import io
import csv
import time # Import time for sleep
from flask import Flask, jsonify, render_template, Response, request
from collections import Counter
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename
import docx
from docx.opc.constants import RELATIONSHIP_TYPE
import PyPDF2
from PyPDF2.generic import IndirectObject, DictionaryObject # Import necessary PyPDF2 types
from langchain_community.llms import Ollama
from langchain_community.embeddings import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from chromadb.api import models
from chromadb import HttpClient
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser


app = Flask(__name__)

# --- Constants ---
ATTACK_ENTERPRISE_URL = "https://raw.githubusercontent.com/mitre/cti/master/enterprise-attack/enterprise-attack.json"
SEARCH_ROOT_DIR = "Searches"
ATTACK_DATA_CACHE = None
SEARCH_DATA_CACHE = None
TECHNIQUE_TO_GROUP_MAP_CACHE = None

# Define a common User-Agent header to mimic a web browser
DEFAULT_HEADERS = {
    'User-Agent': 'Mozilla/50 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
}

# Add ALLOWED_EXTENSIONS and helper function here
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'json', 'csv', 'html'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# --- AI Model Configuration ---
OLLAMA_BASE_URL = os.getenv('OLLAMA_BASE_URL', 'http://ollama:11434')
LLM_MODEL = "qwen2.5vl:3b" # Changed from phi4-reasoning:latest to qwen2.5vl:3b
EMBEDDING_MODEL = "nomic-embed-text:latest"
CHROMA_HOST = os.getenv('CHROMA_HOST', 'chromadb')
CHROMA_PORT = os.getenv('CHROMA_PORT', '8000')
CHROMA_COLLECTION_NAME = "mitre_attack_knowledge" # Collection for ATT&CK data
CHROMA_USER_DATA_COLLECTION_NAME = "user_uploaded_content" # New collection for user data

# Debugging: Print the OLLAMA_BASE_URL as seen by the Flask app
print(f"DEBUG: OLLAMA_BASE_URL (from env) is: {OLLAMA_BASE_URL}", flush=True)

# Global variables for the ChromaDB vectorstores
attack_vectorstore = None
user_data_vectorstore = None

# Global variables for Ollama instances (initialized with retry)
llm = None
embeddings = None

# --- Service Readiness Helper ---
def wait_for_service(service_name, host, port, max_retries=30, delay_seconds=5):
    """
    Waits for a service to be available by attempting to connect.
    Returns True if successful, False otherwise.
    """
    print(f"Waiting for {service_name} at {host}:{port}...", flush=True)
    for i in range(max_retries):
        try:
            if service_name == "Ollama":
                # Directly use OLLAMA_BASE_URL which is expected to have the scheme
                # Append /api/tags for the health check endpoint
                ollama_health_check_url = f"{OLLAMA_BASE_URL}/api/tags"
                print(f"  DEBUG: Attempting Ollama API check at: {ollama_health_check_url}", flush=True)
                response = requests.get(ollama_health_check_url, timeout=5)
                response.raise_for_status()
                print(f"{service_name} is ready! Status: {response.status_code}", flush=True)
                return True
            elif service_name == "ChromaDB":
                # For ChromaDB, try connecting the client
                # Use the full URL for the heartbeat endpoint
                chroma_heartbeat_url = f"http://{host}:{port}/api/v1/heartbeat"
                print(f"  Attempting ChromaDB heartbeat at: {chroma_heartbeat_url}", flush=True)
                response = requests.get(chroma_heartbeat_url, timeout=5)
                response.raise_for_status()
                print(f"{service_name} is ready! Status: {response.status_code}", flush=True)
                return True
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout, requests.exceptions.RequestException) as e:
            print(f"Attempt {i+1}/{max_retries}: {service_name} not ready yet ({e}). Retrying in {delay_seconds}s...", flush=True)
            if isinstance(e, requests.exceptions.HTTPError):
                print(f"  HTTP Error Status: {e.response.status_code}, Response: {e.response.text}", flush=True)
            time.sleep(delay_seconds)
        except Exception as e:
            print(f"Attempt {i+1}/{max_retries}: An unexpected error occurred while waiting for {service_name}: {e}. Retrying in {delay_seconds}s...", flush=True)
            time.sleep(delay_seconds)
    print(f"Failed to connect to {service_name} after {max_retries} retries.", flush=True)
    return False

def initialize_ai_components():
    """Initializes Ollama LLM, Embeddings, and ChromaDB clients with retry logic."""
    global llm, embeddings, attack_vectorstore, user_data_vectorstore

    print("AI_INIT: Starting AI components initialization...", flush=True)
    print(f"AI_INIT: LLM_MODEL is: {LLM_MODEL}, EMBEDDING_MODEL is: {EMBEDDING_MODEL}", flush=True)

    # 1. Initialize Ollama LLM and Embeddings
    print("AI_INIT: Attempting to initialize Ollama LLM and Embeddings...", flush=True)
    ollama_host_match = re.match(r"https?://([^:]+):?(\d+)?", OLLAMA_BASE_URL)
    ollama_host = ollama_host_match.group(1) if ollama_host_match else 'ollama'
    ollama_port = ollama_host_match.group(2) if ollama_host_match and ollama_host_match.group(2) else '11434'

    if not wait_for_service("Ollama", ollama_host, ollama_port):
        print("AI_INIT: Ollama service could not be reached. AI features will be unavailable.", flush=True)
        llm = None
        embeddings = None
        attack_vectorstore = None
        user_data_vectorstore = None
        return # Cannot proceed if Ollama is down

    try:
        llm = Ollama(model=LLM_MODEL, base_url=OLLAMA_BASE_URL)
        # Removed all extra parameters from OllamaEmbeddings as they caused validation errors
        embeddings = OllamaEmbeddings(
            model=EMBEDDING_MODEL,
            base_url=OLLAMA_BASE_URL
        )
        print("AI_INIT: Ollama LLM and Embeddings initialized successfully.", flush=True)
        print(f"AI_INIT: LLM is {llm is not None}, Embeddings is {embeddings is not None}", flush=True) # Confirm state
    except Exception as e:
        print(f"AI_INIT: ERROR: Failed to initialize Ollama LLM/Embeddings: {e}. AI features will be unavailable.", flush=True)
        llm = None
        embeddings = None
        attack_vectorstore = None
        user_data_vectorstore = None
        return # Cannot proceed if embeddings are not available

    # 2. Initialize ChromaDB vectorstores ONLY if embeddings are available
    if embeddings:
        print("AI_INIT: Embeddings are available. Proceeding with ChromaDB initialization.", flush=True)
        if not wait_for_service("ChromaDB", CHROMA_HOST, CHROMA_PORT):
            print("AI_INIT: ChromaDB service could not be reached. Vector store features will be unavailable.", flush=True)
            attack_vectorstore = None
            user_data_vectorstore = None
            return

        try:
            print("AI_INIT: Calling load_attack_data_to_chroma() to initialize attack_vectorstore.", flush=True)
            load_attack_data_to_chroma() # This depends on 'embeddings'
            print("AI_INIT: Calling get_user_data_vectorstore() to initialize user_data_vectorstore.", flush=True)
            get_user_data_vectorstore() # This also depends on 'embeddings'
            print("AI_INIT: ChromaDB vectorstores initialization complete.", flush=True)
        except Exception as e:
            print(f"AI_INIT: ERROR: Failed to initialize ChromaDB vectorstores: {e}. Vector store features will be unavailable.", flush=True)
            attack_vectorstore = None
            user_data_vectorstore = None
    else:
        print("AI_INIT: Embeddings were NOT initialized, skipping ChromaDB vectorstore setup.", flush=True)

    print("AI_INIT: AI components initialization attempt complete.", flush=True)


# --- Caching and Data Fetching ---

def get_attack_data():
    """Fetches and caches the MITRE ATT&CK Enterprise data."""
    global ATTACK_DATA_CACHE
    if ATTACK_DATA_CACHE is None:
        print("Fetching latest ATT&CK data from MITRE...", flush=True)
        try:
            response = requests.get(ATTACK_ENTERPRISE_URL, headers=DEFAULT_HEADERS)
            response.raise_for_status()
            ATTACK_DATA_CACHE = response.json()
            print("ATT&CK data fetched and cached successfully.", flush=True)
        except requests.RequestException as e:
            print(f"Error fetching ATT&CK data: {e}. Returning None.", flush=True)
            return None
    return ATTACK_DATA_CACHE

def _process_searches(all_searches):
    """Processes search data to identify covered techniques and map searches to techniques."""
    print("Starting to process search data...", flush=True)
    covered_techniques = set()
    tech_to_search_map = {}

    for search_name, details in all_searches.items():
        tags_field = details.get('tags', [])
        tag_list = _get_tag_list(tags_field)
        
        for tag in tag_list:
            if tag.strip().startswith('mitre.T'):
                technique_id = tag.strip().split('mitre.')[1]
                covered_techniques.add(technique_id)
                if technique_id not in tech_to_search_map:
                    tech_to_search_map[technique_id] = []
                
                search_info = {
                    "name": search_name,
                    "query": details.get("query", ""),
                    "os": _determine_os(details.get("query", "")),
                    "source_file": details.get("source_file", ""),
                    "html_url": details.get("html_url")
                }
                tech_to_search_map[technique_id].append(search_info)
    print("Search data processing complete.", flush=True)
    return covered_techniques, tech_to_search_map


def get_all_searches():
    """
    Acts as a cached data provider for all search YAML files.
    It orchestrates fetching from Git or local disk.
    """
    global SEARCH_DATA_CACHE
    if SEARCH_DATA_CACHE is None:
        print("Attempting to get all searches...", flush=True)
        searches_dir_url = os.getenv('SEARCHES_REPO_DIR_URL')
        if searches_dir_url:
            SEARCH_DATA_CACHE = _get_searches_from_git(searches_dir_url)
        else:
            SEARCH_DATA_CACHE = _get_searches_from_local()
        print(f"Finished getting all searches. Found {len(SEARCH_DATA_CACHE)} searches.", flush=True)
    return SEARCH_DATA_CACHE

def _get_searches_from_local():
    """Reads all YAML files from the local SEARCH_ROOT_DIR, prioritizing 'searches.yaml'."""
    print(f"SEARCHES_REPO_DIR_URL not set. Falling back to local directory: '{SEARCH_ROOT_DIR}'", flush=True)
    all_searches = {}
    if not os.path.isdir(SEARCH_ROOT_DIR):
        print(f"Warning: Local directory '{SEARCH_ROOT_DIR}' not found. Returning empty searches.", flush=True)
        return {}

    # Prioritize loading from 'searches.yaml' if it exists
    searches_yaml_path = os.path.join(SEARCH_ROOT_DIR, 'searches.yaml')
    if os.path.exists(searches_yaml_path):
        print(f"  - Loading primary search file: {searches_yaml_path}", flush=True)
        try:
            with open(searches_yaml_path, 'r', encoding='utf-8') as f:
                yaml_content = yaml.load(f, Loader=yaml.BaseLoader)
                if isinstance(yaml_content, dict):
                    for search_name, details in yaml_content.items():
                        if isinstance(details, dict):
                            details['source_file'] = searches_yaml_path
                            details['html_url'] = None
                            all_searches[search_name] = details
            print(f"  - Successfully loaded {len(all_searches)} searches from {searches_yaml_path}.", flush=True)
            return all_searches # Return immediately if searches.yaml is found and loaded
        except Exception as e:
            print(f"!!!!!!!!!!!!!! ERROR: Could not process primary search file {searches_yaml_path}. Falling back to recursive search. !!!!!!!!!!!!!!", flush=True)
            print(f"  Reason: {e}", flush=True)

    # Fallback to recursive search if searches.yaml not found or failed to load
    print(f"  - No primary searches.yaml found or failed to load. Recursively searching '{SEARCH_ROOT_DIR}' for all .yml/.yaml files.", flush=True)
    for root, _, files in os.walk(SEARCH_ROOT_DIR):
        for filename in files:
            if filename.endswith(('.yml', '.yaml')):
                filepath = os.path.join(root, filename)
                print(f"  - Loading local file: {filepath}", flush=True)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        yaml_content = yaml.load(f, Loader=yaml.BaseLoader)
                        if isinstance(yaml_content, dict):
                            for search_name, details in yaml_content.items():
                                if isinstance(details, dict):
                                    details['source_file'] = filepath
                                    details['html_url'] = None
                                    all_searches[search_name] = details
                except Exception as e:
                    print(f"!!!!!!!!!!!!!! ERROR: Could not process file {filepath}. Skipping. !!!!!!!!!!!!!!", flush=True)
                    print(f"  Reason: {e}", flush=True)

    return all_searches

def _get_searches_from_git(repo_url):
    """Recursively reads all YAML files from a directory in a Git repo."""
    print(f"Fetching searches from Git repo: {repo_url}", flush=True)
    all_searches = {}
    
    def fetch_dir(api_url):
        try:
            dir_response = requests.get(api_url, headers=DEFAULT_HEADERS) # Added headers
            dir_response.raise_for_status()
            items = dir_response.json()
            if not isinstance(items, list):
                print(f"Warning: Could not retrieve directory contents from {api_url}. Message: {items.get('message')}. Returning.", flush=True)
                return

            for item in items:
                if item['type'] == 'file' and item['name'].endswith(('.yml', '.yaml')):
                    try:
                        file_url = item['download_url']
                        print(f"  - Loading file: {item['path']}", flush=True)
                        file_response = requests.get(file_url, headers=DEFAULT_HEADERS) # Added headers
                        file_response.raise_for_status()
                        yaml_content = yaml.load(file_response.text, Loader=yaml.BaseLoader)
                        if isinstance(yaml_content, dict):
                            for search_name, details in yaml_content.items():
                                if isinstance(details, dict):
                                    details['source_file'] = item['path']
                                    details['html_url'] = item['html_url']
                                    all_searches[search_name] = details
                    except Exception as e:
                        print(f"!!!!!!!!!!!!!! ERROR: Could not process Git file {item.get('path', 'unknown')}. Skipping. !!!!!!!!!!!!!!", flush=True)
                        print(f"  Reason: {e}", flush=True)

                elif item['type'] == 'dir':
                    fetch_dir(item['url'])
        except requests.RequestException as e:
            print(f"Error during GitHub API request to {api_url}: {e}. Returning.", flush=True)

    match = re.match(r"https://github\.com/([^/]+)/([^/]+)/tree/([^/]+)/?(.*)", repo_url)
    if not match:
        print(f"URL format not recognized for Git repo: {repo_url}. Please use format: https://github.com/user/repo/tree/branch. Returning empty searches.", flush=True)
        return {}
    owner, repo, branch, _ = match.groups()
    initial_api_url = f"https://api.github.com/repos/{owner}/{repo}/contents/{SEARCH_ROOT_DIR}?ref={branch}"
    
    fetch_dir(initial_api_url)
    return all_searches

# --- Utility Functions ---

def _get_tag_list(tags_field):
    """Safely converts a YAML tags field to a list of strings."""
    if not tags_field:
        return []
    if isinstance(tags_field, list):
        return [str(tag) for tag in tags_field]
    if isinstance(tags_field, str):
        return tags_field.split(',')
    return []

def _determine_os(query):
    """Determines the target OS based on keywords in the search query."""
    if not isinstance(query, str):
        return "Unknown"
        
    q = query.lower()
    if any(k in q for k in ['aws', 'azure', 'gcp', 'kubernetes', 's3', 'az-']):
        return "Cloud"
    if any(k in q for k in ['macos', 'osx', 'launchctl', 'osascript', '.plist']):
        return "macOS"
    if any(k in q for k in ['windows', 'win', '.exe', 'powershell', 'cmd', 'hkey', 'dll', 'nt authority']):
        return "Windows"
    if any(k in q for k in ['linux', 'bash', 'sudo', 'chmod', '/var/log', '/etc/']):
        return "Linux"
    return "Cross-Platform"

def _build_technique_to_groups_map():
    """
    Builds and caches a map from MITRE ATT&CK technique IDs to a list of threat groups,
    where each group includes both its name and MITRE ID (G-code).
    Ensures G-codes are prioritized and duplicates are avoided.
    """
    global TECHNIQUE_TO_GROUP_MAP_CACHE
    if TECHNIQUE_TO_GROUP_MAP_CACHE is not None:
        return TECHNIQUE_TO_GROUP_MAP_CACHE
    
    print("Building technique-to-group mapping...", flush=True)
    attack_data = get_attack_data()
    if not attack_data:
        print("Could not get ATT&CK data to build technique-to-group mapping. Returning empty map.", flush=True)
        return {}
        
    tech_to_group_temp = {} # Use a temporary dict to build the map
    stix_objects = {obj['id']: obj for obj in attack_data['objects']}

    for obj in attack_data['objects']:
        if obj.get("revoked", False): continue # Skip revoked objects
        if obj.get('type') == 'relationship' and obj.get('relationship_type') == 'uses':
            group_stix_id = obj.get('source_ref')
            tech_stix_id = obj.get('target_ref')

            group = stix_objects.get(group_stix_id)
            tech = stix_objects.get(tech_stix_id)
            
            if group and tech and group.get('type') == 'intrusion-set' and tech.get('type') == 'attack-pattern':
                mitre_tech_id = next((ref['external_id'] for ref in tech.get('external_references', []) if ref.get('source_name') == 'mitre-attack' and 'external_id' in ref), None)
                
                group_name = group.get('name', 'Unknown Group')
                
                # Collect all MITRE external IDs for this group
                mitre_external_ids = [
                    ref['external_id'] for ref in group.get('external_references', [])
                    if ref.get('source_name') == 'mitre-attack' and 'external_id' in ref
                ]

                g_code_id = None
                other_mitre_id = None

                # Find the G-code first
                for ext_id in mitre_external_ids:
                    if ext_id.startswith('G'):
                        g_code_id = ext_id
                        break
                
                # If no G-code, find any other mitre-attack ID
                if g_code_id is None:
                    other_mitre_id = next((ext_id for ext_id in mitre_external_ids), None)

                # Determine the final ID for the map: G-code > Other MITRE ID > Group Name
                final_group_id_for_map = g_code_id or other_mitre_id or group_name
                
                print(f"DEBUG_MAP_DETAIL: Processing Group '{group_name}' (STIX ID: {group_stix_id})", flush=True)
                print(f"DEBUG_MAP_DETAIL:   External References: {group.get('external_references', [])}", flush=True)
                print(f"DEBUG_MAP_DETAIL:   Derived G-code: {g_code_id}, Other MITRE ID: {other_mitre_id}", flush=True)
                print(f"DEBUG_MAP_DETAIL:   Final ID for map: '{final_group_id_for_map}'", flush=True)
                
                if mitre_tech_id and final_group_id_for_map:
                    if mitre_tech_id not in tech_to_group_temp:
                        tech_to_group_temp[mitre_tech_id] = {} # Use a dictionary keyed by group name for this technique
                    
                    current_groups_dict_for_tech = tech_to_group_temp[mitre_tech_id]
                    
                    # If the group name is already in the dict
                    if group_name in current_groups_dict_for_tech:
                        existing_group_info = current_groups_dict_for_tech[group_name]
                        # If existing ID is NOT a G-code, but the new ID IS a G-code, update it
                        if not existing_group_info['id'].startswith('G') and final_group_id_for_map.startswith('G'):
                            current_groups_dict_for_tech[group_name] = {'name': group_name, 'id': final_group_id_for_map}
                            print(f"DEBUG_MAP_ADD: UPDATED group '{group_name}' for technique {mitre_tech_id} to G-code: {final_group_id_for_map}", flush=True)
                        # Else, if the new ID is a G-code and it's different from the existing G-code (shouldn't happen often)
                        elif final_group_id_for_map.startswith('G') and existing_group_info['id'].startswith('G') and existing_group_info['id'] != final_group_id_for_map:
                            print(f"DEBUG_MAP_ADD: WARNING: Found conflicting G-codes for '{group_name}' for technique {mitre_tech_id}. Keeping existing: {existing_group_info['id']}, new: {final_group_id_for_map}", flush=True)
                    else:
                        # Add new group
                        current_groups_dict_for_tech[group_name] = {'name': group_name, 'id': final_group_id_for_map}
                        print(f"DEBUG_MAP_ADD: Added group '{group_name}' (ID: {final_group_id_for_map}) for technique {mitre_tech_id}", flush=True)
    
    # Convert the inner dictionaries back to lists of values
    TECHNIQUE_TO_GROUP_MAP_CACHE = {
        tech_id: list(groups_dict.values())
        for tech_id, groups_dict in tech_to_group_temp.items()
    }

    print("Technique-to-group mapping built.", flush=True)
    return TECHNIQUE_TO_GROUP_MAP_CACHE

def _get_processed_matrix_data():
    """Helper function to generate the matrix data structure."""
    print("Starting to process matrix data...", flush=True)
    all_searches = get_all_searches()
    attack_json = get_attack_data()

    if not attack_json:
        print("Error: Could not retrieve ATT&CK data from MITRE for matrix processing. Returning error.", flush=True)
        return {"error": "Could not retrieve ATT&CK data from MITRE."}
    if not all_searches:
        print("Warning: No search data found for matrix processing.", flush=True)

    covered_techniques, tech_to_search_map = _process_searches(all_searches)
    
    objects = attack_json.get("objects", [])
    tactics, techniques_map = {}, {}
    
    print(f"Processing {len(objects)} STIX objects for matrix...", flush=True)
    for o in objects:
        if o.get("revoked", False): continue
        object_type = o.get("type")
        ext_refs = o.get("external_references", [])
        
        if object_type == "x-mitre-tactic" and ext_refs:
            ext_ref = ext_refs[0]
            tactics[o['x_mitre_shortname']] = {"name": o["name"], "id": ext_ref["external_id"], "url": ext_ref["url"], "techniques": []}
        elif object_type == "attack-pattern" and ext_refs:
            ext_ref = ext_refs[0]
            tech_id = ext_ref["external_id"]
            is_subtechnique = o.get('x_mitre_is_subtechnique', False)
            parent_id = tech_id.split('.')[0] if is_subtechnique else None
            techniques_map[tech_id] = {
                "id": tech_id, "name": o["name"], "url": ext_ref["url"],
                "covered": tech_id in covered_techniques, "searches": tech_to_search_map.get(tech_id, []),
                "kill_chain_phases": o.get("kill_chain_phases", []),
                "is_subtechnique": is_subtechnique, "parent_id": parent_id, "subtechniques": []
            }

    print("Mapping techniques to tactics and sub-techniques to techniques for matrix...", flush=True)
    for tech_id, tech_info in techniques_map.items():
        if tech_info["is_subtechnique"]:
            if (parent_id := tech_info.get("parent_id")) and (parent := techniques_map.get(parent_id)):
                parent["subtechniques"].append(tech_info)
        else:
            for phase in tech_info.get("kill_chain_phases", []):
                if phase.get("kill_chain_name") == "mitre-attack":
                    if tactic := tactics.get(phase["phase_name"]):
                        tactic["techniques"].append(tech_info)
    
    print("Sorting tactics and techniques for matrix...", flush=True)
    sorted_tactics = sorted(tactics.values(), key=lambda x: int(x['id'].replace('TA','')))
    for tactic in sorted_tactics:
        tactic['techniques'].sort(key=lambda x: x['name'])
        for tech in tactic['techniques']:
            # Corrected syntax error here: removed the extra 'x:'
            tech['subtechniques'].sort(key=lambda x: x['name'])
    
    print(f"Matrix data processing complete. Number of tactics: {len(sorted_tactics)}.", flush=True)
    # Print a snippet of the processed data for verification
    if sorted_tactics:
        print(f"First tactic in processed data: {sorted_tactics[0]['name']} (ID: {sorted_tactics[0]['id']})", flush=True)
        if sorted_tactics[0]['techniques']:
            print(f"First technique in first tactic: {sorted_tactics[0]['techniques'][0]['name']} (ID: {sorted_tactics[0]['techniques'][0]['id']})", flush=True)
    else:
        print("Processed matrix data is empty.", flush=True)

    return sorted_tactics

def load_attack_data_to_chroma():
    """Loads processed MITRE ATT&CK data into ChromaDB for AI retrieval."""
    global attack_vectorstore
    if embeddings is None:
        print("CHROMA_ATTACK: Embeddings not initialized, skipping ATT&CK data loading to ChromaDB.", flush=True)
        attack_vectorstore = None
        return

    print("CHROMA_ATTACK: Attempting to load ATT&CK data into ChromaDB...", flush=True)

    # Fetch ATT&CK data
    matrix_data = _get_processed_matrix_data()
    if isinstance(matrix_data, dict) and matrix_data.get("error"):
        print(f"CHROMA_ATTACK: Error getting matrix data for ChromaDB: {matrix_data['error']}", flush=True)
        attack_vectorstore = None
        return

    documents = []
    # Add techniques and sub-techniques as documents
    for tactic in matrix_data:
        # Add tactic itself as a document
        documents.append({
            "id": tactic['id'],
            "content": f"MITRE ATT&CK Tactic: {tactic['name']}. Description URL: {tactic['url']}",
            "metadata": {"type": "tactic", "name": tactic['name'], "id": tactic['id'], "url": tactic['url']}
        })
        for tech in tactic['techniques']:
            content = f"MITRE ATT&CK Technique: {tech['name']} (ID: {tech['id']}). Description URL: {tech['url']}."
            if tech['covered']:
                content += f" This technique is covered by {len(tech['searches'])} searches: {', '.join([s['name'] for s in tech['searches']])}."
            documents.append({
                "id": tech['id'],
                "content": content,
                "metadata": {"type": "technique", "name": tech['name'], "id": tech['id'], "url": tech['url'], "covered": tech['covered']}
            })
            for subtech in tech['subtechniques']:
                sub_content = f"MITRE ATT&CK Sub-technique: {subtech['name']} (ID: {subtech['id']}). Parent Technique: {tech['name']} ({tech['id']}). Description URL: {subtech['url']}."
                if subtech['covered']:
                    sub_content += f" This sub-technique is covered by {len(subtech['searches'])} searches: {', '.join([s['name'] for s in subtech['searches']])}."
                documents.append({
                    "id": subtech['id'],
                    "content": sub_content,
                    "metadata": {"type": "subtechnique", "name": subtech['name'], "id": subtech['id'], "url": subtech['url'], "covered": subtech['covered'], "parent_id": tech['id']}
                })

    # Convert to LangChain Document objects
    lc_documents = [Document(page_content=doc["content"], metadata=doc["metadata"]) for doc in documents]

    # Use a text splitter to break down larger documents if necessary
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=150, chunk_overlap=50)
    split_docs = text_splitter.split_documents(lc_documents)

    # Initialize ChromaDB vectorstore for ATT&CK data
    try:
        print(f"CHROMA_ATTACK: Initializing client for ATT&CK data at host={CHROMA_HOST}, port={CHROMA_PORT}...", flush=True)
        chroma_client = HttpClient(host=CHROMA_HOST, port=int(CHROMA_PORT))
        print("CHROMA_ATTACK: Client for ATT&CK data created. Attempting heartbeat...", flush=True)
        chroma_client.heartbeat() # Test connection
        print("CHROMA_ATTACK: Client heartbeat successful for ATT&CK data.", flush=True)

        # Check if collection exists and delete/recreate for a fresh start (useful during development)
        try:
            print(f"CHROMA_ATTACK: Attempting to delete existing collection '{CHROMA_COLLECTION_NAME}'...", flush=True)
            chroma_client.delete_collection(name=CHROMA_COLLECTION_NAME)
            print(f"CHROMA_ATTACK: Deleted existing collection: {CHROMA_COLLECTION_NAME}", flush=True)
        except Exception as e:
            print(f"CHROMA_ATTACK: Collection '{CHROMA_COLLECTION_NAME}' did not exist or could not be deleted. Proceeding to create/get. ({e})", flush=True)
        
        print(f"CHROMA_ATTACK: Adding {len(split_docs)} documents to collection '{CHROMA_COLLECTION_NAME}'...", flush=True)
        attack_vectorstore = Chroma.from_documents(
            documents=split_docs,
            embedding=embeddings,
            client=chroma_client,
            collection_name=CHROMA_COLLECTION_NAME
        )
        print(f"CHROMA_ATTACK: ATT&CK data loaded into collection '{CHROMA_COLLECTION_NAME}'.", flush=True)
        print(f"CHROMA_ATTACK: Number of items in '{CHROMA_COLLECTION_NAME}' collection: {attack_vectorstore._collection.count()}", flush=True) # Verify count
    except Exception as e:
        print(f"CHROMA_ATTACK: ERROR: Error initializing or loading ATT&CK data to ChromaDB: {e}", flush=True)
        attack_vectorstore = None # Ensure it's None if initialization fails

def get_user_data_vectorstore():
    """Initializes or retrieves the ChromaDB vectorstore for user-uploaded content."""
    global user_data_vectorstore
    print("CHROMA_USER_DATA: Attempting to get user data vectorstore...", flush=True)
    print(f"CHROMA_USER_DATA: Current 'embeddings' object state: {embeddings is not None}", flush=True)

    if embeddings is None:
        print("CHROMA_USER_DATA: WARNING: Embeddings not initialized, cannot initialize user data vectorstore. Returning None.", flush=True)
        user_data_vectorstore = None
        return None

    # Check if user_data_vectorstore is None OR if its internal embedding function is not set
    # This addresses potential issues where the object exists but isn't fully configured for embedding
    if user_data_vectorstore is None or not hasattr(user_data_vectorstore, '_embedding_function') or user_data_vectorstore._embedding_function is None:
        try:
            print(f"CHROMA_USER_DATA: Initializing client for user data at host={CHROMA_HOST}, port={CHROMA_PORT}...", flush=True)
            chroma_client = HttpClient(host=CHROMA_HOST, port=int(CHROMA_PORT))
            print("CHROMA_USER_DATA: Client for user data created. Attempting heartbeat...", flush=True)
            chroma_client.heartbeat() # Test connection
            print("CHROMA_USER_DATA: Client heartbeat successful for user data.", flush=True)

            # Get or create the collection directly to ensure it exists
            # This is a more robust way to ensure the collection is ready
            collection = chroma_client.get_or_create_collection(name=CHROMA_USER_DATA_COLLECTION_NAME)
            
            # Re-initialize Chroma object, explicitly passing the collection and embedding function
            # Removed _collection argument as it caused an error in previous logs
            user_data_vectorstore = Chroma(
                client=chroma_client,
                collection_name=CHROMA_USER_DATA_COLLECTION_NAME,
                embedding_function=embeddings
            )
            print(f"CHROMA_USER_DATA: User data vectorstore initialized for collection: {CHROMA_USER_DATA_COLLECTION_NAME}", flush=True)
            print(f"CHROMA_USER_DATA: Number of items in '{CHROMA_USER_DATA_COLLECTION_NAME}' collection: {user_data_vectorstore._collection.count()}", flush=True) # Verify count
            print(f"CHROMA_USER_DATA: user_data_vectorstore._embedding_function is: {user_data_vectorstore._embedding_function is not None}", flush=True) # New check
        except Exception as e:
            print(f"CHROMA_USER_DATA: ERROR: Error initializing user data vectorstore: {e}", flush=True)
            user_data_vectorstore = None
    else:
        print("CHROMA_USER_DATA: User data vectorstore already initialized and has embedding function. Returning existing instance.", flush=True)
        print(f"CHROMA_USER_DATA: user_data_vectorstore._embedding_function is: {user_data_vectorstore._embedding_function is not None}", flush=True) # New check
    return user_data_vectorstore

# --- API Endpoints ---
@app.route('/')
def index():
    """Render the main HTML page."""
    return render_template('index.html')

@app.route('/api/analyze_document', methods=['POST'])
def analyze_document():
    """
    Analyzes multiple uploaded documents for T-codes and their associated groups,
    and embeds their content into the user_data_vectorstore.
    """
    print("DOC_ANALYSIS: Received analyze document request (multiple files).", flush=True)

    # Use getlist to handle multiple files uploaded under the same field name 'files'
    files = request.files.getlist('files')

    if not files:
        print("DOC_ANALYSIS: No files part in the request.", flush=True)
        return jsonify({"error": "No files selected for upload"}), 400

    # Ensure AI components are initialized before processing
    if llm is None or embeddings is None:
        print("DOC_ANALYSIS: LLM or Embeddings not initialized. Attempting re-initialization...", flush=True)
        initialize_ai_components() # Attempt to re-initialize
        if llm is None or embeddings is None:
            print("DOC_ANALYSIS: AI models still not initialized after re-attempt. Returning 503.", flush=True)
            return jsonify({"error": "AI models not initialized. Please check backend services."}), 503

    all_document_results = []
    # Initialize text splitter once for all documents
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    for file in files:
        document_result = {
            "filename": secure_filename(file.filename) if file.filename else "unknown_file",
            "status": "processed",
            "techniques": [],
            "urls": [],
            "error": None
        }
        
        if not file.filename:
            document_result["status"] = "failed"
            document_result["error"] = "No filename provided."
            all_document_results.append(document_result)
            print(f"DOC_ANALYSIS: Skipping file due to no filename: {file}", flush=True)
            continue

        filename = secure_filename(file.filename)
        print(f"DOC_ANALYSIS: Processing uploaded file: {filename} ({file.content_type}).", flush=True)

        text_content = ""
        urls = []

        try:
            if filename.endswith('.pdf'):
                reader = PyPDF2.PdfReader(file.stream)
                for page in reader.pages:
                    text_content += page.extract_text() or ""
                    if '/Annots' in page:
                        for annot in page['/Annots']:
                            # Safely resolve IndirectObject and check if it's a dictionary-like object
                            resolved_annot = annot.get_object()
                            if isinstance(resolved_annot, (DictionaryObject, dict)):
                                subtype = resolved_annot.get('/Subtype')
                                if subtype == '/Link':
                                    uri_dict = resolved_annot.get('/A')
                                    if isinstance(uri_dict, (DictionaryObject, dict)) and uri_dict.get('/S') == '/URI':
                                        urls.append(uri_dict.get('/URI'))
            elif filename.endswith('.docx'):
                doc = docx.Document(file.stream)
                text_parts = [para.text for para in doc.paragraphs]
                text_content = "\n".join(text_parts)
                for rel in doc.part.rels.values():
                    if rel.reltype == RELATIONSHIP_TYPE.HYPERLINK:
                        if rel.target_ref not in urls:
                            urls.append(rel.target_ref)
            else:
                document_result["status"] = "failed"
                document_result["error"] = "Unsupported file type. Only PDF and DOCX are supported."
                print(f"DOC_ANALYSIS: Unsupported file type for {filename}: {file.content_type}", flush=True)
                all_document_results.append(document_result)
                continue # Skip to next file

            print(f"DOC_ANALYSIS: Extracted text length for {filename}: {len(text_content)}", flush=True)

            # Analyze T-codes
            t_codes = sorted(list(set(re.findall(r'(T\d{4}(?:\.\d{3})?)', text_content))))
            tech_to_group_map = _build_technique_to_groups_map() # Ensure this is cached or efficient
            techniques_with_groups = []
            for t_code in t_codes:
                # Retrieve the full group info (name and ID) from the map
                groups_info = tech_to_group_map.get(t_code, [])
                techniques_with_groups.append({
                    "id": t_code,
                    "groups": groups_info # This now contains list of {'name': ..., 'id': ...}
                })
            document_result["techniques"] = techniques_with_groups
            document_result["urls"] = urls

            # --- Embed document content into ChromaDB ---
            if text_content.strip(): # Only embed if there's actual text content
                user_vs = get_user_data_vectorstore()
                
                # Explicitly check the condition's evaluation
                condition_met = (user_vs is not None and embeddings is not None and hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None)
                print(f"DOC_ANALYSIS: Embedding condition evaluation: {condition_met} (user_vs={user_vs is not None}, embeddings={embeddings is not None}, user_vs._embedding_function={hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None})", flush=True)

                if condition_met:
                    print(f"DOC_ANALYSIS: Entering embedding block for {filename}!", flush=True) # Confirm entry
                    try:
                        doc_to_embed = Document(
                            page_content=text_content,
                            metadata={"source": filename, "type": "document_upload"}
                        )
                        print(f"DOC_ANALYSIS: Splitting document {filename} into chunks...", flush=True)
                        split_docs = text_splitter.split_documents([doc_to_embed])
                        print(f"DOC_ANALYSIS: Created {len(split_docs)} chunks for {filename}.", flush=True)
                        
                        print(f"DOC_ANALYSIS: Embedding and adding {len(split_docs)} chunks to ChromaDB for {filename}...", flush=True)
                        user_vs.add_documents(split_docs)
                        print(f"DOC_ANALYSIS: Successfully added {len(split_docs)} chunks to user_data vectorstore for {filename}.", flush=True)
                    except Exception as e:
                        document_result["status"] = "failed"
                        document_result["error"] = f"Error embedding content: {e}"
                        print(f"DOC_ANALYSIS: Error embedding document '{filename}' into ChromaDB: {e}", flush=True)
                else: # THIS ELSE BLOCK IS STILL BEING HIT DESPITE TRUE PRINTS
                    print(f"DOC_ANALYSIS: ERROR: Embedding condition NOT met for {filename}. user_vs is {user_vs is not None}, embeddings is {embeddings is not None}, user_vs._embedding_function is {hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None}", flush=True)
                    document_result["status"] = "skipped_embedding"
                    document_result["error"] = "User data vectorstore or embeddings not available, skipping document embedding."
                    print(f"DOC_ANALYSIS: User data vectorstore or embeddings not available for {filename}, skipping document embedding.", flush=True)
            else:
                print(f"DOC_ANALYSIS: No extractable text content from {filename}, skipping embedding.", flush=True)
            
        except Exception as e:
            document_result["status"] = "failed"
            document_result["error"] = f"Failed to process file: {e}"
            print(f"DOC_ANALYSIS: Error processing document {filename}: {e}", flush=True)

        all_document_results.append(document_result)

    return jsonify({"results": all_document_results})

@app.route('/api/analyze_url', methods=['POST'])
def analyze_url():
    """Analyzes a given URL for T-codes and their associated groups, and embeds its content."""
    print("URL_ANALYSIS: Received analyze URL request.", flush=True)
    data = request.get_json()
    url_to_analyze = data.get('url')

    if not url_to_analyze:
        print("URL_ANALYSIS: No URL provided.", flush=True)
        return jsonify({"error": "No URL provided"}), 400

    # Ensure AI components are initialized before processing
    if llm is None or embeddings is None:
        print("URL_ANALYSIS: LLM or Embeddings not initialized. Attempting re-initialization...", flush=True)
        initialize_ai_components() # Attempt to re-initialize
        if llm is None or embeddings is None:
            print("URL_ANALYSIS: AI models still not initialized after re-attempt. Returning 503.", flush=True)
            return jsonify({"error": "AI models not initialized. Please check backend services."}), 503

    text_content = ""
    analysis_status = "processed" # Initialize status for this single URL analysis
    analysis_error = None
    word_count = 0 # Initialize word count

    try:
        print(f"URL_ANALYSIS: Attempting to fetch content from URL: {url_to_analyze}", flush=True)
        response = requests.get(url_to_analyze, headers=DEFAULT_HEADERS, timeout=10)
        response.raise_for_status()
        print(f"URL_ANALYSIS: Successfully fetched content from {url_to_analyze}. Status: {response.status_code}", flush=True)
        
        # Parse HTML content
        soup = BeautifulSoup(response.text, 'html.parser')
        # Extract text from common tags, excluding script and style
        for script_or_style in soup(['script', 'style']):
            script_or_style.extract() # Remove these tags
        text_content = soup.get_text(separator='\n', strip=True)
        
        # Calculate word count
        word_count = len(text_content.split())
        print(f"URL_ANALYSIS: Word count for {url_to_analyze}: {word_count}", flush=True)

    except requests.exceptions.HTTPError as e:
        analysis_status = "failed"
        analysis_error = f"Failed to fetch content from URL: HTTP Status {e.response.status_code}"
        print(f"URL_ANALYSIS: Failed to fetch content from {url_to_analyze}: HTTP Status {e.response.status_code}", flush=True)
    except requests.exceptions.RequestException as e:
        analysis_status = "failed"
        analysis_error = f"Error fetching content from URL: {e}"
        print(f"URL_ANALYSIS: Error fetching content from URL {url_to_analyze}: {e}", flush=True)
    except Exception as e:
        analysis_status = "failed"
        analysis_error = f"An unexpected error occurred during URL content extraction: {e}"
        print(f"URL_ANALYSIS: An unexpected error occurred during URL content extraction: {e}", flush=True)

    print(f"URL_ANALYSIS: Extracted text length: {len(text_content)}", flush=True)

    t_codes = []
    techniques_with_groups = []
    if analysis_status == "processed": # Only try to extract T-codes if content fetch was successful
        t_codes = sorted(list(set(re.findall(r'(T\d{4}(?:\.\d{3})?)', text_content))))
        tech_to_group_map = _build_technique_to_groups_map()
        for t_code in t_codes:
            # Retrieve the full group info (name and ID) from the map
            groups_info = tech_to_group_map.get(t_code, [])
            techniques_with_groups.append({
                "id": t_code,
                "groups": groups_info # This now contains list of {'name': ..., 'id': ...}
            })

    # --- Embed URL content into ChromaDB ---
    if text_content.strip() and analysis_status == "processed": # Only embed if there's actual text content AND no prior error
        user_vs = get_user_data_vectorstore()
        
        # Explicitly check the condition's evaluation
        condition_met = (user_vs is not None and embeddings is not None and hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None)
        print(f"URL_ANALYSIS: Embedding condition evaluation: {condition_met} (user_vs={user_vs is not None}, embeddings={embeddings is not None}, user_vs._embedding_function={hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None})", flush=True)

        if condition_met:
            print(f"URL_ANALYSIS: Entering embedding block for {url_to_analyze}!", flush=True) # Confirm entry
            try:
                # Create a LangChain Document from the extracted text
                doc_to_embed = Document(
                    page_content=text_content,
                    metadata={"source": url_to_analyze, "type": "url_analysis"}
                )
                # Split the document into chunks for embedding
                print("URL_ANALYSIS: Splitting document into chunks...", flush=True)
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                split_docs = text_splitter.split_documents([doc_to_embed])
                print(f"URL_ANALYSIS: Created {len(split_docs)} chunks.", flush=True)
                
                print("URL_ANALYSIS: Embedding and adding chunks to ChromaDB...", flush=True)
                user_vs.add_documents(split_docs)
                print(f"URL_ANALYSIS: Successfully added {len(split_docs)} chunks to user_data vectorstore.", flush=True)
            except Exception as e:
                analysis_status = "failed"
                analysis_error = f"Error embedding content: {e}"
                print(f"URL_ANALYSIS: Error embedding URL content into ChromaDB: {e}", flush=True)
        else:
            analysis_status = "skipped_embedding"
            analysis_error = "User data vectorstore or embeddings not available, skipping URL content embedding."
            print(f"URL_ANALYSIS: ERROR: Embedding condition NOT met for {url_to_analyze}. user_vs is {user_vs is not None}, embeddings is {embeddings is not None}, user_vs._embedding_function is {hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None}", flush=True)
    elif not text_content.strip():
        print(f"URL_ANALYSIS: No extractable text content from {url_to_analyze}, skipping embedding.", flush=True)
    # --- End Embedding ---

    return jsonify({
        "source_url": url_to_analyze,
        "status": analysis_status,
        "error": analysis_error,
        "techniques": techniques_with_groups,
        "urls": [url_to_analyze], # Return the analyzed URL
        "word_count": word_count # Added word count
    })

@app.route('/api/ai_query', methods=['POST'])
def ai_query():
    """Handles AI queries using the RAG chain."""
    print("AI_QUERY: Received request.", flush=True)
    data = request.get_json()
    user_query = data.get('query')
    use_user_data = data.get('use_user_data', False) # New parameter

    if not user_query:
        print("AI_QUERY: No query provided.", flush=True)
        return jsonify({"error": "No query provided"}), 400

    if llm is None or embeddings is None:
        print("AI_QUERY: LLM or Embeddings not initialized. Attempting re-initialization...", flush=True)
        initialize_ai_components() # Attempt to re-initialize
        if llm is None or embeddings is None:
            print("AI_QUERY: AI models still not initialized after re-attempt. Returning 503.", flush=True)
            return jsonify({"error": "AI models not initialized. Please check backend services."}), 503
    
    # Determine which retriever to use
    retriever_to_use = None
    if use_user_data:
        user_vs = get_user_data_vectorstore()
        if user_vs:
            retriever_to_use = user_vs.as_retriever()
            print("AI_QUERY: Using user_data_vectorstore for retrieval.", flush=True)
        else:
            print("AI_QUERY: User data requested but user_data_vectorstore not available. Falling back to ATT&CK data.", flush=True)
    
    if retriever_to_use is None and attack_vectorstore:
        retriever_to_use = attack_vectorstore.as_retriever()
        print("AI_QUERY: Using attack_vectorstore for retrieval (either by default or as fallback).", flush=True)
    
    if retriever_to_use is None:
        print("AI_QUERY: No active knowledge base (attack_vectorstore or user_data_vectorstore) initialized. Returning 503.", flush=True)
        return jsonify({"error": "No active knowledge base initialized. Please analyze documents or ensure ATT&CK data is loaded."}), 503

    try:
        # Define the prompt template for the LLM
        # UPDATED PROMPT TEMPLATE FOR SPECIALIZED AI
        template = """You are a highly specialized AI assistant for cybersecurity, focusing on threat intelligence,
        Indicators of Compromise (IOCs), MITRE ATT&CK framework, malware analysis, and general cyber security issues.

        When responding, consider the following:
        - Summarize documents that are uploaded and provide compared context if multiple documents are relevant.
        - If asked about detection rules, provide examples.
        - If the answer cannot be found in the provided context, politely state that you don't have enough information
          but offer to search the broader MITRE ATT&CK knowledge base or suggest uploading more relevant documents.
	- Speak like you are Will from the movie Good Will Hunting
	- Do you have indicators or other data that is related to eachother

        Context:
        {context}

        Question: {question}
        """
        prompt = ChatPromptTemplate.from_template(template)
        print("AI_QUERY: Prompt template created.", flush=True) # Debugging

        # Build the RAG chain
        rag_chain = (
            {"context": retriever_to_use, "question": RunnablePassthrough()}
            | prompt
            | llm
            | StrOutputParser()
        )
        print("AI_QUERY: RAG chain built.", flush=True) # Debugging

        # Invoke the chain with the user's query
        ai_response = rag_chain.invoke(user_query)
        print(f"AI_QUERY: Received response from LLM (first 100 chars): {ai_response[:100]}...", flush=True) # Debugging
        return jsonify({"response": ai_response})

    except Exception as e:
        print(f"AI_QUERY: Error during AI query processing: {e}", flush=True) # Debugging
        return jsonify({"error": f"Failed to process AI query: {e}. Check backend logs for details."}), 500

@app.route('/api/matrix')
def get_matrix():
    """Returns the processed ATT&CK matrix data."""
    print("MATRIX: Received request for matrix data.", flush=True)
    matrix_data = _get_processed_matrix_data()
    if isinstance(matrix_data, dict) and matrix_data.get("error"):
        print(f"MATRIX: Error returning matrix data: {matrix_data['error']}", flush=True)
        return jsonify(matrix_data), 500
    print("MATRIX: Returning matrix data.", flush=True)
    return jsonify(matrix_data)

@app.route('/api/techniques/<string:technique_id>/groups')
def get_technique_groups(technique_id):
    """
    Returns a list of threat groups (with name and ID) that use the given technique.
    """
    print(f"TECH_GROUPS: Received request for groups for technique ID: {technique_id}", flush=True)
    technique_to_group_map = _build_technique_to_groups_map()
    # The map now stores {'name': ..., 'id': ...} objects, so we return them directly
    groups_info = technique_to_group_map.get(technique_id.upper(), [])
    print(f"TECH_GROUPS: Found {len(groups_info)} groups for {technique_id}.", flush=True)
    return jsonify({"groups": groups_info})

@app.route('/api/technique_usage/<string:group_id>/<string:technique_id>')
def get_technique_usage(group_id, technique_id):
    """
    Fetches and returns procedural usage for a specific technique by a specific group
    from the MITRE ATT&CK CTI data.
    """
    print(f"TECH_USAGE: Received request for usage for group ID/Name '{group_id}' and technique '{technique_id}'", flush=True)
    attack_data = get_attack_data()
    if not attack_data:
        print("TECH_USAGE: ATT&CK data not loaded.", flush=True)
        return jsonify({"error": "ATT&CK data not loaded."}), 500

    group_stix_id = None
    group_name = None

    # Try to find the group by its external_id (G-code) or name
    print(f"TECH_USAGE: Searching for group '{group_id}' in ATT&CK objects...", flush=True)
    for obj in attack_data['objects']:
        if obj['type'] == 'intrusion-set':
            current_group_name = obj.get('name', 'N/A')
            current_external_references = obj.get('external_references', [])
            print(f"TECH_USAGE:   Checking intrusion-set: Name='{current_group_name}', STIX_ID='{obj['id']}', External_Refs={current_external_references}", flush=True)

            # Check by external_id (preferring G-code match)
            for ref in current_external_references:
                if ref.get('source_name') == 'mitre-attack' and ref.get('external_id') == group_id:
                    group_stix_id = obj['id']
                    group_name = obj['name']
                    print(f"TECH_USAGE:   MATCH! Found group '{group_name}' by external_id '{group_id}'. STIX ID: {group_stix_id}", flush=True)
                    break
            if group_stix_id: # If found by external_id, break outer loop
                break
            
            # If not found by external_id, try by name (case-insensitive for robustness)
            if current_group_name.lower() == group_id.lower():
                group_stix_id = obj['id']
                group_name = obj['name']
                print(f"TECH_USAGE:   MATCH! Found group '{current_group_name}' by name (case-insensitive) '{group_id}'. STIX ID: {group_stix_id}", flush=True)
                break # Found by name, break outer loop

    if not group_stix_id:
        print(f"TECH_USAGE: Group with ID/Name '{group_id}' NOT FOUND in ATT&CK data after all checks.", flush=True)
        return jsonify({"error": f"Group with ID/Name '{group_id}' not found in ATT&CK data."}), 404

    technique_stix_id = None
    # Find the STIX ID for the given technique_id (T-code)
    print(f"TECH_USAGE: Searching for technique '{technique_id}' in ATT&CK objects...", flush=True)
    for obj in attack_data['objects']:
        if obj['type'] == 'attack-pattern' and 'external_references' in obj:
            for ref in obj['external_references']:
                if ref.get('source_name') == 'mitre-attack' and ref.get('external_id') == technique_id:
                    technique_stix_id = obj['id']
                    print(f"TECH_USAGE:   MATCH! Found technique '{technique_id}'. STIX ID: {technique_stix_id}", flush=True)
                    break
        if technique_stix_id:
            break

    if not technique_stix_id:
        print(f"TECH_USAGE: Technique with ID {technique_id} NOT FOUND in ATT&CK data.", flush=True)
        return jsonify({"error": f"Technique with ID {technique_id} not found in ATT&CK data."}), 404

    # Find the relationship object that describes the usage
    usage_description = "No specific procedural usage found for this technique by this group in the ATT&CK data."
    
    print(f"TECH_USAGE: Searching for 'uses' relationship between group '{group_stix_id}' and technique '{technique_stix_id}'...", flush=True)
    for obj in attack_data['objects']:
        if obj['type'] == 'relationship' and \
           obj['relationship_type'] == 'uses' and \
           obj['source_ref'] == group_stix_id and \
           obj['target_ref'] == technique_stix_id:
            usage_description = obj.get('description', usage_description)
            print(f"TECH_USAGE: Found usage description for {group_name} and {technique_id}.", flush=True)
            break
    
    # If no specific description, try to find the technique description itself
    if usage_description == "No specific procedural usage found for this technique by this group in the ATT&CK data.":
        technique_obj = next((item for item in attack_data['objects'] if item['id'] == technique_stix_id and item['type'] == 'attack-pattern'), None)
        if technique_obj and technique_obj.get('description'):
            usage_description = f"General description for {technique_id} ({technique_obj['name']}): {technique_obj['description']}"
            print(f"TECH_USAGE: No specific group usage found, falling back to general technique description for {technique_id}.", flush=True)

    return jsonify({
        "group_id": group_id,
        "group_name": group_name,
        "technique_id": technique_id,
        "usage": usage_description
    })

@app.route('/api/missing_tags')
def get_missing_tags():
    """Returns a list of searches that do not have a valid MITRE ATT&CK tag."""
    print("MISSING_TAGS: Received request for missing tags.", flush=True)
    all_searches = get_all_searches()
    if not all_searches:
        print("MISSING_TAGS: No search data available.", flush=True)
        return jsonify({"error": "No search data available."}), 500

    missing_tags_searches = []
    for search_name, details in all_searches.items():
        tags_field = details.get('tags', [])
        tag_list = _get_tag_list(tags_field)
        
        has_mitre_tag = False
        for tag in tag_list:
            if tag.strip().startswith('mitre.T'):
                has_mitre_tag = True
                break
        
        if not has_mitre_tag:
            missing_tags_searches.append({
                "name": search_name,
                "query": details.get("query", "N/A"),
                "source_file": details.get("source_file", "N/A"),
                "html_url": details.get("html_url", None)
            })
    print(f"MISSING_TAGS: Found {len(missing_tags_searches)} searches with missing tags.", flush=True)
    return jsonify(missing_tags_searches)

@app.route('/api/missing_tags/csv')
def export_missing_tags_csv():
    """Exports the list of searches with missing MITRE ATT&CK tags to a CSV file."""
    print("EXPORT_MISSING_TAGS: Received request to export missing tags to CSV.", flush=True)
    missing_tags_data = get_missing_tags().json # Get the JSON data from the route
    if missing_tags_data.get("error"):
        print(f"EXPORT_MISSING_TAGS: Error getting data: {missing_tags_data['error']}", flush=True)
        return missing_tags_data, 500

    si = io.StringIO()
    cw = csv.writer(si)

    # Write headers
    cw.writerow(["Search Name", "Query", "Source File", "HTML URL"])

    # Write data rows
    for search in missing_tags_data:
        cw.writerow([search["name"], search["query"], search["source_file"], search["html_url"]])

    output = si.getvalue()
    print("EXPORT_MISSING_TAGS: CSV generated successfully.", flush=True)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=missing_mitre_tags.csv"}
    )

@app.route('/api/matrix/csv')
def export_matrix_csv():
    """Exports the ATT&CK matrix data to a CSV file."""
    print("EXPORT_MATRIX: Received request to export matrix to CSV.", flush=True)
    matrix_data = _get_processed_matrix_data()
    if isinstance(matrix_data, dict) and matrix_data.get("error"):
        print(f"EXPORT_MATRIX: Error getting data: {matrix_data['error']}", flush=True)
        return jsonify(matrix_data), 500

    si = io.StringIO()
    cw = csv.writer(si)

    # Write headers: Tactic Name, Technique ID, Technique Name, Covered, Searches (comma-separated), URL
    cw.writerow(["Tactic Name", "Technique ID", "Technique Name", "Covered by Search", "Searches", "URL"])

    for tactic in matrix_data:
        for technique in tactic['techniques']:
            searches_str = ", ".join([s['name'] for s in technique['searches']])
            cw.writerow([
                tactic['name'],
                technique['id'],
                technique['name'],
                "Yes" if technique['covered'] else "No",
                searches_str,
                technique['url']
            ])
            for subtech in technique['subtechniques']:
                sub_searches_str = ", ".join([s['name'] for s in subtech['searches']])
                cw.writerow([
                    tactic['name'],
                    subtech['id'],
                    f"  {subtech['name']}", # Indent sub-techniques
                    "Yes" if subtech['covered'] else "No",
                    sub_searches_str,
                    subtech['url']
                ])
    
    output = si.getvalue()
    print("EXPORT_MATRIX: CSV generated successfully.", flush=True)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=mitre_attack_coverage.csv"}
    )

@app.route('/api/os_distribution')
def get_os_distribution():
    """Returns the distribution of searches by operating system."""
    print("OS_DIST: Received request for OS distribution.", flush=True)
    all_searches = get_all_searches()
    if not all_searches:
        print("OS_DIST: No search data available.", flush=True)
        return jsonify({"error": "No search data available."}), 500

    os_counts = Counter()
    for search_name, details in all_searches.items():
        query = details.get("query", "")
        os_name = _determine_os(query)
        os_counts[os_name] += 1
    
    print("OS_DIST: OS distribution calculated.", flush=True)
    return jsonify(dict(os_counts))

@app.route('/api/searches_by_os/<string:os_name>')
def get_searches_by_os(os_name):
    """Returns a list of searches for a specific operating system."""
    print(f"SEARCHES_BY_OS: Received request for searches by OS: {os_name}", flush=True)
    all_searches = get_all_searches()
    if not all_searches:
        print("SEARCHES_BY_OS: No search data available.", flush=True)
        return jsonify({"error": "No search data available."}), 500

    filtered_searches = []
    for search_name, details in all_searches.items():
        query = details.get("query", "")
        if _determine_os(query) == os_name:
            filtered_searches.append({
                "name": search_name,
                "query": query,
                "source_file": details.get("source_file", "N/A"),
                "html_url": details.get("html_url", None)
            })
    print(f"SEARCHES_BY_OS: Found {len(filtered_searches)} searches for OS: {os_name}.", flush=True)
    return jsonify(filtered_searches)

@app.route('/api/searches_by_os/<string:os_name>/csv')
def export_searches_by_os_csv(os_name):
    """Exports searches for a specific OS to a CSV file."""
    print(f"EXPORT_OS_SEARCHES: Received request to export searches for OS: {os_name} to CSV.", flush=True)
    searches_data = get_searches_by_os(os_name).json # Get the JSON data from the route
    if searches_data.get("error"):
        print(f"EXPORT_OS_SEARCHES: Error getting data: {searches_data['error']}", flush=True)
        return searches_data, 500

    si = io.StringIO()
    cw = csv.writer(si)

    # Write headers
    cw.writerow(["Search Name", "Query", "Source File", "HTML URL"])

    # Write data rows
    for search in searches_data:
        cw.writerow([search["name"], search["query"], search["source_file"], search["html_url"]])

    output = si.getvalue()
    print(f"EXPORT_OS_SEARCHES: CSV generated successfully for OS: {os_name}.", flush=True)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment;filename={os_name.lower().replace(' ', '_')}_searches.csv"}
    )

@app.route('/api/analyze_github_repo', methods=['POST'])
def analyze_github_repo():
    """
    Analyzes a GitHub repository (or a specific path within it) for T-codes,
    extracts code snippets, and embeds content into the user_data_vectorstore.
    """
    print("GITHUB_ANALYSIS: Received analyze GitHub repo request.", flush=True)
    data = request.get_json()
    repo_url = data.get('repo_url')

    if not repo_url:
        print("GITHUB_ANALYSIS: No GitHub repository URL provided.", flush=True)
        return jsonify({"error": "No GitHub repository URL provided"}), 400

    # Ensure AI components are initialized before processing
    if llm is None or embeddings is None:
        print("GITHUB_ANALYSIS: LLM or Embeddings not initialized. Attempting re-initialization...", flush=True)
        initialize_ai_components() # Attempt to re-initialize
        if llm is None or embeddings is None:
            print("GITHUB_ANALYSIS: AI models still not initialized after re-attempt. Returning 503.", flush=True)
            return jsonify({"error": "AI models not initialized. Please check backend services."}), 503

    # Regex to parse GitHub URL: https://github.com/owner/repo/tree/branch/path/to/dir
    match = re.match(r"https://github\.com/([^/]+)/([^/]+)(?:/tree/([^/]+)/(.*))?", repo_url)
    if not match:
        print(f"GITHUB_ANALYSIS: Invalid GitHub URL format: {repo_url}", flush=True)
        return jsonify({"error": "Invalid GitHub URL format. Expected: https://github.com/owner/repo or https://github.com/owner/repo/tree/branch/path"}), 400

    owner, repo_name, branch, path = match.groups()
    branch = branch if branch else "main" # Default to 'main' branch
    path = path if path else "" # Default to root of the repo

    api_url = f"https://api.github.com/repos/{owner}/{repo_name}/contents/{path}?ref={branch}"
    
    all_file_results = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    def fetch_github_content(url):
        try:
            print(f"GITHUB_ANALYSIS: Fetching GitHub API content from: {url}", flush=True)
            response = requests.get(url, headers=DEFAULT_HEADERS, timeout=15)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.HTTPError as e:
            print(f"GITHUB_ANALYSIS: HTTP Error fetching GitHub content from {url}: {e.response.status_code} - {e.response.text}", flush=True)
            return {"error": f"GitHub API error: {e.response.status_code} - {e.response.text}"}
        except requests.exceptions.RequestException as e:
            print(f"GITHUB_ANALYSIS: Network Error fetching GitHub content from {url}: {e}", flush=True)
            return {"error": f"Network error fetching GitHub content: {e}"}

    def process_item(item):
        file_result = {
            "file_path": item.get('path', 'N/A'),
            "status": "processed",
            "word_count": 0,
            "techniques": [],
            "code_snippets": [],
            "error": None
        }
        
        if item['type'] == 'file':
            # Skip binary files or very large files
            if item['size'] > 1024 * 1024 * 2: # 2 MB limit
                file_result["status"] = "skipped"
                file_result["error"] = "File too large."
                all_file_results.append(file_result)
                print(f"GITHUB_ANALYSIS: Skipping large file: {item['path']}", flush=True)
                return
            
            # Heuristic to skip common binary/unparseable files
            if any(item['name'].lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.zip', '.tar', '.gz', '.7z', '.rar', '.exe', '.dll', '.so', '.dylib', '.pyc', '.class']):
                file_result["status"] = "skipped"
                file_result["error"] = "Likely binary or unparseable file type."
                all_file_results.append(file_result)
                print(f"GITHUB_ANALYSIS: Skipping binary/unparseable file: {item['path']}", flush=True)
                return

            try:
                print(f"GITHUB_ANALYSIS: Fetching content for file: {item['path']}", flush=True)
                file_content_response = requests.get(item['download_url'], headers=DEFAULT_HEADERS, timeout=10)
                file_content_response.raise_for_status()
                file_content = file_content_response.text

                file_result["word_count"] = len(file_content.split())

                # Extract T-codes
                t_codes = sorted(list(set(re.findall(r'(T\d{4}(?:\.\d{3})?)', file_content))))
                tech_to_group_map = _build_technique_to_groups_map()
                techniques_with_groups = []
                for t_code in t_codes:
                    groups_info = tech_to_group_map.get(t_code, [])
                    techniques_with_groups.append({
                        "id": t_code,
                        "groups": groups_info
                    })
                file_result["techniques"] = techniques_with_groups

                # Extract code snippets (simple heuristic: lines starting with common code keywords)
                code_lines = []
                for line in file_content.splitlines():
                    stripped_line = line.strip()
                    if stripped_line and (
                        stripped_line.startswith(('def ', 'class ', 'import ', 'from ', 'func ', 'function ', 'var ', 'const ', 'let ', 'public ', 'private ', 'protected ', '#include', 'using namespace', 'console.log', 'printf', 'System.out')) or
                        re.match(r'^[a-zA-Z_]\w*\s*\(.*\)\s*{', stripped_line) # function/method definition
                    ):
                        code_lines.append(line)
                        if len(code_lines) > 50: # Limit snippet size
                            break
                if code_lines:
                    file_result["code_snippets"].append({
                        "file_path": item['path'],
                        "code": "\n".join(code_lines)
                    })

                # Embed content into ChromaDB
                if file_content.strip():
                    user_vs = get_user_data_vectorstore()
                    condition_met = (user_vs is not None and embeddings is not None and hasattr(user_vs, '_embedding_function') and user_vs._embedding_function is not None)
                    print(f"GITHUB_ANALYSIS: Embedding condition evaluation for {item['path']}: {condition_met}", flush=True)

                    if condition_met:
                        try:
                            doc_to_embed = Document(
                                page_content=file_content,
                                metadata={"source": item['html_url'], "type": "github_repo_analysis", "file_path": item['path']}
                            )
                            split_docs = text_splitter.split_documents([doc_to_embed])
                            user_vs.add_documents(split_docs)
                            print(f"GITHUB_ANALYSIS: Successfully added {len(split_docs)} chunks for {item['path']} to user_data vectorstore.", flush=True)
                        except Exception as e:
                            file_result["status"] = "failed"
                            file_result["error"] = f"Error embedding content: {e}"
                            print(f"GITHUB_ANALYSIS: Error embedding GitHub file '{item['path']}' into ChromaDB: {e}", flush=True)
                    else:
                        file_result["status"] = "skipped_embedding"
                        file_result["error"] = "User data vectorstore or embeddings not available, skipping content embedding."
                        print(f"GITHUB_ANALYSIS: User data vectorstore or embeddings not available for {item['path']}, skipping embedding.", flush=True)
                else:
                    print(f"GITHUB_ANALYSIS: No extractable text content from {item['path']}, skipping embedding.", flush=True)

            except requests.exceptions.RequestException as e:
                file_result["status"] = "failed"
                file_result["error"] = f"Failed to download file content: {e}"
                print(f"GITHUB_ANALYSIS: Failed to download file content {item['path']}: {e}", flush=True)
            except Exception as e:
                file_result["status"] = "failed"
                file_result["error"] = f"An unexpected error occurred during file processing: {e}"
                print(f"GITHUB_ANALYSIS: An unexpected error occurred processing file {item['path']}: {e}", flush=True)
        
        all_file_results.append(file_result)

    repo_contents = fetch_github_content(api_url)

    if "error" in repo_contents:
        return jsonify({"error": repo_contents["error"]}), 500

    # If the URL points directly to a file, repo_contents will be a dict, not a list
    if isinstance(repo_contents, dict) and repo_contents.get('type') == 'file':
        process_item(repo_contents)
    elif isinstance(repo_contents, list):
        for item in repo_contents:
            if item['type'] == 'file':
                process_item(item)
            elif item['type'] == 'dir':
                # Recursively fetch directory contents
                dir_contents = fetch_github_content(item['url'])
                if isinstance(dir_contents, list):
                    for sub_item in dir_contents:
                        process_item(sub_item)
                elif "error" in dir_contents:
                    all_file_results.append({
                        "file_path": item.get('path', 'N/A'),
                        "status": "failed",
                        "error": f"Could not list directory contents: {dir_contents['error']}"
                    })
    else:
        return jsonify({"error": "Unexpected response from GitHub API. Not a file or directory listing."}), 500

    return jsonify({"results": all_file_results})

@app.route('/api/user_data', methods=['GET'])
def get_user_data():
    global user_data_vectorstore
    if user_data_vectorstore is None:
        return jsonify({"error": "User data vector store not initialized. Upload documents first."}), 503

    try:
        # Access the underlying Chroma collection to retrieve all documents
        # Note: For very large collections, consider implementing pagination
        # or using a more robust iteration method.
        # 'limit=None' typically means no limit for get(), effectively fetching all.
        all_data = user_data_vectorstore._collection.get(
            ids=[],          # Empty list of IDs means no specific ID filter
            where={},        # Empty dict means no metadata filter
            limit=None,      # Fetch all available documents
            include=['documents', 'metadatas'] # Specify what to include in the response
        )

        formatted_data = []
        if 'documents' in all_data and 'metadatas' in all_data:
            for i in range(len(all_data['documents'])):
                formatted_data.append({
                    "content": all_data['documents'][i],
                    "metadata": all_data['metadatas'][i]
                })

        return jsonify({"data": formatted_data})

    except Exception as e:
        print(f"Error retrieving user data from vector store: {e}", flush=True)
        return jsonify({"error": f"Failed to retrieve user data: {str(e)}"}), 500

@app.route('/api/fetch_navigator_url', methods=['POST'])
def fetch_navigator_url():
    """
    Proxy endpoint to fetch external MITRE ATT&CK Navigator layer JSON files.
    This avoids CORS issues when the browser tries to fetch directly from external domains.
    """
    data = request.get_json()
    if not data or 'url' not in data:
        return jsonify({"error": "No URL provided in request body."}), 400
    
    url = data['url']
    try:
        print(f"FETCH_NAVIGATOR: Attempting to fetch layer from: {url}", flush=True)
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        # Verify it's JSON
        try:
            layer_data = response.json()
            return jsonify(layer_data)
        except ValueError:
            return jsonify({"error": "The URL did not return a valid JSON response."}), 400

    except requests.exceptions.HTTPError as e:
        print(f"ERROR: Failed to fetch Navigator URL {url}: {e}", flush=True)
        return jsonify({"error": f"External server returned an error: {e.response.status_code}"}), e.response.status_code
    except requests.exceptions.RequestException as e:
        print(f"ERROR: Request error fetching Navigator URL {url}: {e}", flush=True)
        return jsonify({"error": f"Failed to connect to the external URL: {str(e)}"}), 502
    except Exception as e:
        print(f"ERROR: Unexpected error fetching Navigator URL {url}: {e}", flush=True)
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

if __name__ == '__main__':
    # Initialize ATT&CK data and searches on startup
    print("APP_START: Starting app initialization...", flush=True)
    get_attack_data()
    get_all_searches()
    _build_technique_to_groups_map() # Pre-build this map
    initialize_ai_components() # This will handle Ollama and ChromaDB initialization
    
    app.run(debug=True, host='0.0.0.0')